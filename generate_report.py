#!/usr/bin/env python3
"""Gera relatório DOCX do projeto HealthFit."""
import sys
sys.path.insert(0, "/Users/berglimma/Documents/HealhtFit/.docx_deps")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUTPUT = "/Users/berglimma/Documents/HealhtFit/HealthFit/Relatorio_HealthFit_Funcionalidades_e_CodeReview.docx"


def set_doc_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.color.rgb = RGBColor(0x1A, 0x5C, 0x3A)


def add_title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("HealthFit\n")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x3A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Relatório Técnico de Funcionalidades e Code Review")
    sr.font.size = Pt(16)
    sr.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nVersão do documento: 1.0\nData: {date.today().strftime('%d/%m/%Y')}\n")
    meta.add_run("Plataforma: iOS + watchOS (Swift / SwiftUI)\n")
    meta.add_run("Desenvolvedor: BERG / LUAN - BLSwift Solutions LTDA\n")

    doc.add_page_break()


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_function(doc, name, objective, details=None):
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    p.add_run(f" — Objetivo: {objective}")
    if details:
        d = doc.add_paragraph(details)
        d.paragraph_format.left_indent = Inches(0.25)


def build_report():
    doc = Document()
    set_doc_styles(doc)
    add_title_page(doc)

    # 1. RESUMO
    doc.add_heading("1. Resumo Executivo", level=1)
    doc.add_paragraph(
        "O HealthFit é um aplicativo de fitness desenvolvido em SwiftUI para iPhone com "
        "companion para Apple Watch. O sistema cobre treinos de musculação, cardio, meditação "
        "guiada, nutrição, bem-estar (sono e hidratação), integração com HealthKit, relatórios "
        "semanais, envio de e-mail ao personal trainer, notificações inteligentes e contagem de "
        "repetições via câmera (Vision). A persistência é local (UserDefaults e arquivos), sem "
        "backend remoto. Este documento detalha cada módulo, suas funções, objetivos e apresenta "
        "um code review técnico completo."
    )

    # 2. ARQUITETURA
    doc.add_heading("2. Arquitetura do Sistema", level=1)
    doc.add_paragraph(
        "Padrão MVVM adaptado ao SwiftUI: Views observam serviços @MainActor via @EnvironmentObject. "
        "O ponto de composição é HealthFitApp.swift, que instancia AuthService, WorkoutStore, "
        "HealthKitManager, MealPlanService, RestTimerService, WatchConnectivityManager, "
        "WeeklyReportService e DailyWellnessService."
    )
    add_bullet(doc, "Navegação: NavigationStack com navigationDestination tipado")
    add_bullet(doc, "Modais: fullScreenCover para treinos ativos; sheet para check-in, e-mail e relatórios")
    add_bullet(doc, "Persistência: UserDefaults (JSON) + Documents (foto de perfil)")
    add_bullet(doc, "Watch: WCSession com protocolo de mensagens por chave action")
    add_bullet(doc, "Helpers estáticos: WeeklyProgressAnalyzer, WorkoutReportBuilder, MotivationMessages")

    # 3. FUNCIONALIDADES
    doc.add_heading("3. Funcionalidades Implementadas (Detalhamento)", level=1)

    modules = [
        ("3.1 Autenticação e Perfil", "Services/AuthService.swift · Views/Auth/ · Views/Profile/", [
            ("login(email:password:)", "Autenticar usuário localmente após validação mínima de campos.", "Cria UserProfile a partir do e-mail; persiste em UserDefaults; agenda notificações."),
            ("register(name:email:password:biotype:goal:)", "Cadastrar novo atleta com biotipo e objetivo.", "Inicia sessão local e gera plano alimentar se necessário."),
            ("logout()", "Encerrar sessão e limpar dados locais sensíveis.", "Remove usuário, foto e cancela lembretes."),
            ("updateProfile(_:)", "Atualizar dados do perfil (peso, objetivo, personal trainer).", "Recalcula metas calóricas e hidratação."),
            ("updateProfileImage(_:)", "Salvar/remover foto de perfil redimensionada.", "JPEG em Documents com nome derivado do e-mail."),
            ("UserProfile (modelo)", "Representar atleta e calcular BMR, TDEE, IMC e meta calórica.", "Fórmula Mifflin-St Jeor; ajustes por biotipo e objetivo."),
            ("ProfileView", "Interface de configuração completa.", "Biotipo, objetivo, personal, sono, água, cronômetro de descanso, integrações."),
        ]),
        ("3.2 Treinos de Musculação", "Services/WorkoutStore.swift · Views/Workout/", [
            ("addWorkoutSheet / updateWorkoutSheet / deleteWorkoutSheet", "CRUD de fichas de treino.", "Persistência em healthfit_workout_sheets."),
            ("startSession(for:)", "Iniciar sessão de musculação.", "Cria WorkoutSession, registros por exercício e timer."),
            ("markExerciseCompleted(at:)", "Marcar exercício/série como concluído.", "Avança índice e atualiza contadores."),
            ("addHeartRateSample / updateCalories", "Atualizar métricas em tempo real.", "Recebe dados do Apple Watch."),
            ("endSession()", "Finalizar treino e salvar histórico.", "Dispara lembrete de inatividade de 48h."),
            ("WorkoutListView / WorkoutDetailView", "Listar e detalhar treinos.", "4 treinos modelo A–D com 8–11 exercícios cada."),
            ("ActiveWorkoutView", "Tela de treino ativo.", "Séries, descanso, overlay de timer, sync com Watch, auto-finish."),
            ("CreateWorkoutView", "Criar ficha personalizada.", "Nome, descrição e lista de exercícios."),
            ("WorkoutSummaryView", "Resumo pós-treino.", "Estatísticas e envio de relatório por e-mail."),
            ("RestTimerOverlay + RestTimerService", "Cronômetro de descanso entre séries.", "Alerta de overtime; sync com Watch; haptics."),
        ]),
        ("3.3 Cardio", "Models/CardioModels.swift · ActiveCardioView.swift", [
            ("CardioExercise.catalog", "Catálogo de 11 exercícios (corrida, bike, escalada, etc.).", "Calorias/minuto estimadas por modalidade."),
            ("CardioIntensity", "Três intensidades: Baixa, Média, Alta.", "Duração sugerida 50/40/30 min e multiplicador calórico."),
            ("CardioSetupView", "Configurar e iniciar sessão cardio.", "Seleção de intensidade e estimativa de calorias."),
            ("ActiveCardioView", "Sessão cardio com timer circular.", "Sync Apple Watch, gravação HealthKit, resumo final."),
            ("startCardioOnWatch / syncCardioProgress", "Sincronizar cardio com relógio.", "Cronômetro e meta no Watch."),
        ]),
        ("3.4 Meditação", "Models/MeditationModels.swift · ActiveMeditationView.swift", [
            ("MeditationTopic.catalog", "7 tópicos com 10 prompts guiados cada.", "Respiração, relaxamento, gratidão, foco, ansiedade, sono, recuperação."),
            ("MeditationDuration", "Durações 5, 10, 15 ou 20 minutos.", "Sessão com prompts rotativos automáticos."),
            ("MeditationSetupView", "Preview dos tópicos e seleção de duração.", "Inicia sessão registrada no histórico."),
            ("ActiveMeditationView", "Timer + orientações em tela.", "Finalização automática ao completar tempo."),
            ("startMeditationSession(config:)", "Registrar meditação como sessão.", "Conta para histórico e lembrete de inatividade."),
        ]),
        ("3.5 Apple Watch", "WatchConnectivityManager.swift · HealthFitWatch/", [
            ("startWorkoutOnWatch / stopWorkoutOnWatch", "Iniciar/encerrar treino no relógio.", "Fallback com dados simulados se offline."),
            ("syncWorkoutProgress", "Enviar cronômetro e exercício atual.", "Atualização a cada segundo."),
            ("sendRestTimerStart / sendRestTimerStop", "Sincronizar descanso.", "Timer visual no Watch."),
            ("WatchContentView", "UI do relógio.", "Cronômetro grande, BPM, kcal, tela de descanso."),
            ("WatchWorkoutManager", "Gerenciar estado e HealthKit no Watch.", "Recebe 12+ tipos de mensagem do iPhone."),
        ]),
        ("3.6 HealthKit e Dashboard", "HealthKitManager.swift · DashboardView.swift", [
            ("requestAuthorization()", "Solicitar permissões de leitura/escrita.", "Passos, calorias, FC, treinos."),
            ("fetchWeeklyMetrics()", "Métricas dos últimos 7 dias.", "Alimenta gráficos do dashboard."),
            ("saveWorkout(...)", "Gravar treino/cardio no Health.", "HKWorkoutBuilder com duração, calorias e FC."),
            ("DashboardView", "Tela inicial.", "Métricas do dia, banner relatório semanal, gráficos, treinos recentes."),
            ("HealthChartsView", "Gráfico de barras Swift Charts.", "Passos, calorias, FC repouso, minutos de treino."),
        ]),
        ("3.7 Relatório Semanal", "WeeklyProgressAnalyzer.swift · WeeklyReportView.swift", [
            ("buildReport(sessions:goal:)", "Gerar relatório dos últimos 7 dias.", "Compara com semana anterior."),
            ("calculateScore", "Pontuação 0–100.", "Frequência, dias ativos, conclusão, minutos, objetivo."),
            ("buildImprovements", "Sugestões personalizadas.", "Frequência, cardio, descanso, volume, objetivo."),
            ("WeeklyReportService", "Disponibilidade a cada 7 dias.", "Badge NOVO no dashboard."),
            ("WeeklyReportView", "UI completa do relatório.", "Score, tendências, gráfico diário, destaques."),
        ]),
        ("3.8 E-mail ao Personal", "WorkoutReportBuilder.swift · MailComposeView.swift", [
            ("emailSubject / emailBody", "Montar relatório em texto.", "Duração, exercícios, descanso, calorias, FC."),
            ("MailComposeView", "Abrir compositor nativo do iOS.", "MFMailComposeViewController com fallback mailto."),
            ("WorkoutSummaryView.sendReportToTrainer", "Fluxo de envio pós-treino.", "Feedback 'E-mail enviado' após confirmação."),
        ]),
        ("3.9 Notificações", "NotificationService.swift · MotivationMessages.swift", [
            ("scheduleDailyMotivationNotifications", "14 dias de motivação às 8h.", "Sincronizado com Apple Watch."),
            ("deliverWorkoutStartNotification / deliverWorkoutEndNotification", "Alertas de início e fim de treino.", "Mensagens personalizadas PT-BR."),
            ("deliverRestOvertimeNotification", "Alerta de descanso prolongado.", "iPhone + Watch."),
            ("refreshWorkoutInactivityReminder", "Lembrete após 48h sem treino.", "Agendamento e entrega com anti-duplicata."),
            ("scheduleInactivityReminderOnWatch", "Espelhar lembrete no relógio.", "UNCalendarNotificationTrigger."),
        ]),
        ("3.10 Bem-estar (Sono e Hidratação)", "DailyWellnessService.swift · DailyWellnessCheckInView.swift", [
            ("configure(for:) / checkInOnAppOpen()", "Exibir check-in diário.", "Solicita horas de sono ao abrir o app."),
            ("SleepAssessment.evaluate(hours:)", "Classificar qualidade do sono.", "<5h irregular; 5–6h precisa mais; 7–9h ideal."),
            ("logSleep(hours:)", "Registrar sono do dia.", "Persistência por usuário e data."),
            ("recommendedDailyWaterML", "Meta de água: 35 ml × peso (kg).", "Exibida no check-in e perfil."),
            ("updateWaterIntake / addWater", "Registrar consumo de água.", "Barra de progresso e mensagens de status."),
            ("ProfileView wellnessSection", "Edição contínua no perfil.", "Slider de sono e stepper de copos."),
        ]),
        ("3.11 Nutrição", "MealPlanService.swift · Views/Nutrition/", [
            ("generatePlan(for:)", "Gerar cardápio semanal.", "Baseado em TDEE, objetivo e biotipo."),
            ("generateShoppingList()", "Lista de compras por categoria.", "Proteínas, carboidratos, hortifruti, etc."),
            ("MealPlanView", "Visualizar plano e macros.", "Edição de peso/altura com recálculo."),
            ("ShoppingListView", "Checklist de compras.", "Marcar itens como comprados."),
        ]),
        ("3.12 Vision / Câmera", "VisionWorkoutService.swift · VisionWorkoutView.swift", [
            ("startDetection / stopDetection", "Pipeline AVCapture + Vision.", "Detecção de pose corporal."),
            ("detectRepetition(hipY:)", "Contador de repetições.", "Baseado em movimento vertical do quadril."),
            ("processBodyPose(_:)", "Verificação de postura.", "Alinhamento ombros/quadril."),
            ("VisionWorkoutView", "UI com overlay de câmera.", "Reps, postura e feedback visual."),
        ]),
    ]

    for title, path, functions in modules:
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Arquivos principais: {path}")
        for name, objective, details in functions:
            add_function(doc, name, objective, details)

    # 4. MODELOS DE DADOS
    doc.add_heading("4. Modelos de Dados Principais", level=1)
    models = [
        ("UserProfile", "Atleta: nome, e-mail, biotipo, objetivo, peso, altura, personal trainer."),
        ("WorkoutSheet / Exercise", "Ficha de treino com exercícios, séries, reps, descanso e grupo muscular."),
        ("WorkoutSession / ExerciseSessionRecord", "Sessão executada com tempos, FC, calorias e conclusão."),
        ("CardioExercise / CardioWorkoutConfig", "Modalidade cardio + intensidade + meta de duração."),
        ("MeditationTopic / MeditationWorkoutConfig", "Tópico de meditação + duração da sessão."),
        ("DailyWellnessEntry / SleepAssessment", "Registro diário de sono e água com classificação."),
        ("WeeklyProgressReport / WeekStats", "Agregados semanais para relatório analítico."),
        ("Meal / DailyMealPlan / ShoppingItem", "Estruturas do módulo nutricional."),
    ]
    for name, desc in models:
        add_bullet(doc, f"{desc}", bold_prefix=f"{name}: ")

    # 5. PERSISTÊNCIA
    doc.add_heading("5. Persistência e Chaves de Armazenamento", level=1)
    keys = [
        "healthfit_current_user — perfil do usuário (AuthService)",
        "healthfit_workout_sheets / healthfit_session_history — treinos (WorkoutStore)",
        "healthfit_meal_plan / healthfit_shopping_list — nutrição (MealPlanService)",
        "healthfit_wellness_{email}_today — sono e água do dia (DailyWellnessService)",
        "healthfit_last_workout_completed_at — base do lembrete 48h (NotificationService)",
        "healthfit_last_weekly_report_viewed — ciclo do relatório semanal",
        "profile_{email}.jpg — foto de perfil em Documents",
    ]
    for k in keys:
        add_bullet(doc, k)

    # 6. CODE REVIEW
    doc.add_page_break()
    doc.add_heading("6. Code Review Detalhado", level=1)

    doc.add_heading("6.1 Pontos Fortes", level=2)
    strengths = [
        "Cobertura funcional ampla para um app fitness: treino, cardio, meditação, nutrição e bem-estar.",
        "UI consistente com tema escuro, componentes reutilizáveis (PrimaryButtonStyle, MetricBadge, cardStyle).",
        "Integração real com HealthKit, Apple Watch e notificações locais.",
        "Relatório semanal com análise comparativa e sugestões baseadas no objetivo do usuário.",
        "Separação razoável entre Models, Services e Views.",
        "Analyzers puros (WeeklyProgressAnalyzer, WorkoutReportBuilder) facilitam manutenção.",
        "Suporte parcial a iPad com adaptiveContentWidth e layout split no treino ativo.",
        "Localização em português brasileiro em toda a experiência do usuário.",
    ]
    for s in strengths:
        add_bullet(doc, s)

    doc.add_heading("6.2 Problemas Críticos e Altos", level=2)
    critical = [
        ("Autenticação mock (Alta)", "AuthService aceita qualquer e-mail/senha sem verificação real. Senha não é armazenada nem validada em logins subsequentes. Inadequado para produção."),
        ("Dados de saúde simulados (Alta)", "HealthKitManager.loadMockData() e fallbacks com Double.random em Watch/HealthKit podem exibir métricas falsas como reais no dashboard."),
        ("PII sem criptografia (Média-Alta)", "UserProfile completo em UserDefaults sem proteção. E-mail do personal e dados físicos expostos em armazenamento não criptografado."),
        ("Status HealthKit enganoso (Média)", "isAuthorized = true após requestAuthorization mesmo se usuário negar permissões granulares."),
        ("Cardio finishCardio inconsistência (Média)", "ActiveCardioView monta exerciseRecords localmente, mas endSession() sobrescreve com exerciseRecords do store não atualizados durante cardio."),
    ]
    for title, desc in critical:
        p = doc.add_paragraph()
        r = p.add_run(title + ": ")
        r.bold = True
        p.add_run(desc)

    doc.add_heading("6.3 Problemas Médios e Baixos", level=2)
    medium = [
        "Ausência total de testes unitários e UI tests.",
        "Sem camada de rede, API ou sincronização em nuvede.",
        "WorkoutStore concentra CRUD, FSM de sessão, timers e migração de samples.",
        "Múltiplos Timer/Timers.publish sem gestão centralizada (risco em background).",
        "VisionWorkoutService isolado — repetições não integram com WorkoutStore.",
        "Meditação não sincroniza com Apple Watch.",
        "isCardioSession usa heurística frágil (prefixo 'cardio' no título).",
        "BiotypeThemes.swift é código morto (duplicata não referenciada).",
        "MealPlanView edita peso/altura localmente; ProfileView só exibe (possível divergência).",
        "Shopping list não soma quantidades de ingredientes duplicados.",
        "Tratamento de erros silencioso em HealthKit save e agendamento de notificações.",
        "Conteúdo hardcoded (treinos, refeições, prompts) dificulta manutenção e i18n.",
    ]
    for m in medium:
        add_bullet(doc, m)

    doc.add_heading("6.4 Segurança e Privacidade", level=2)
    security = [
        "Permissões declaradas: Câmera (Vision), HealthKit (leitura/escrita), Fotos (perfil).",
        "Entitlement HealthKit presente em HealthFit.entitlements.",
        "Recomendação: migrar credenciais e tokens para Keychain; perfil sensível com Data Protection.",
        "Recomendação: autenticação real (Sign in with Apple, Firebase, ou API própria).",
        "Recomendação: não exibir dados simulados quando HealthKit negado — mostrar estado vazio explícito.",
    ]
    for s in security:
        add_bullet(doc, s)

    doc.add_heading("6.5 Qualidade de Código e Manutenibilidade", level=2)
    quality = [
        ("Coesão", "Serviços bem nomeados, porém WorkoutStore e NotificationService cresceram demais — candidatos a refatoração."),
        ("Acoplamento", "Views acessam singletons (NotificationService.shared) além de EnvironmentObject — dificulta testes."),
        ("Concorrência", "Uso correto de @MainActor na maioria dos serviços; Watch delegates usam Task { @MainActor } adequadamente."),
        ("SwiftUI", "Bom uso de navigationDestination, fullScreenCover e @Published; alguns if let dentro de sheet já corrigidos no mail."),
        ("Nomenclatura", "Consistente em português na UI e inglês no código — padrão aceitável para apps BR."),
        ("Documentação", "Poucos comentários; código em geral autoexplicativo, mas falta README técnico."),
    ]
    for title, desc in quality:
        add_function(doc, title, desc)

    doc.add_heading("6.6 Recomendações Prioritárias", level=2)
    recs = [
        "P0: Remover ou isolar dados mock de produção; indicar claramente quando métricas são estimadas.",
        "P0: Implementar autenticação real ou documentar explicitamente como protótipo.",
        "P1: Adicionar testes para WeeklyProgressAnalyzer, SleepAssessment, WorkoutReportBuilder e RestTimerService.",
        "P1: Corrigir fluxo de exerciseRecords no cardio e na meditação.",
        "P1: Unificar edição de peso/altura entre ProfileView e MealPlanView.",
        "P2: Extrair WorkoutSessionManager de WorkoutStore.",
        "P2: Integrar Vision com ActiveWorkoutView para contagem automática de reps.",
        "P2: Adicionar Localizable.strings para internacionalização.",
        "P3: Backend para sync de treinos, relatórios e backup de perfil.",
        "P3: Widget iOS e complicações watchOS para cronômetro e meta de água.",
    ]
    for i, r in enumerate(recs, 1):
        add_bullet(doc, r)

    # 7. MATRIZ
    doc.add_heading("7. Matriz de Funcionalidades por Módulo", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Módulo"
    hdr[1].text = "Status"
    hdr[2].text = "Watch"
    hdr[3].text = "HealthKit"
    rows_data = [
        ("Musculação", "Completo", "Sim", "Sim"),
        ("Cardio", "Completo", "Sim", "Sim"),
        ("Meditação", "Completo", "Não", "Não"),
        ("Nutrição", "Completo", "Não", "Não"),
        ("Sono/Água", "Completo", "Não", "Não"),
        ("Relatório Semanal", "Completo", "Não", "Não"),
        ("E-mail Personal", "Completo", "Não", "Não"),
        ("Notificações", "Completo", "Sim", "Não"),
        ("Vision/Câmera", "Parcial", "Não", "Não"),
        ("Autenticação", "Mock local", "Não", "Não"),
    ]
    for mod, status, watch, hk in rows_data:
        row = table.add_row().cells
        row[0].text = mod
        row[1].text = status
        row[2].text = watch
        row[3].text = hk

    # 8. CONCLUSÃO
    doc.add_heading("8. Conclusão", level=1)
    doc.add_paragraph(
        "O HealthFit apresenta um conjunto robusto de funcionalidades para um aplicativo de fitness "
        "mobile, com destaque para a integração iPhone + Apple Watch, variedade de modalidades de "
        "treino (musculação, cardio e meditação), relatórios analíticos e cuidados com bem-estar "
        "(sono, água e notificações de hábito). O código segue convenções SwiftUI modernas e é "
        "adequado como protótipo funcional ou MVP demonstrativo."
    )
    doc.add_paragraph(
        "Para evolução a produção, as prioridades são: autenticação real, eliminação de dados "
        "simulados enganosos, testes automatizados, refatoração de serviços monolíticos e "
        "fortalecimento da segurança de dados locais. Com esses ajustes, a base arquitetural "
        "existente suporta bem a escalabilidade de novas features."
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— Fim do Relatório —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Relatório gerado: {OUTPUT}")


if __name__ == "__main__":
    build_report()
